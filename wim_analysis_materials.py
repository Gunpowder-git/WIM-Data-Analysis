"""WIMData.csv 作业分析脚本（尽量贴合课堂讲义写法）

特点：
- 读取：使用 csv.DictReader + with open（讲义第52页）
- 编码：可选用 chardet 自动检测（讲义第49页）
- 绘图：使用 matplotlib 画饼图/折线图（讲义第90、102页）

用法（命令行 / Spyder 都可）：
1) 把本脚本与 WIMData.csv 放在同一目录
2) 运行：
   python wim_analysis_materials.py
   或指定文件：
   python wim_analysis_materials.py WIMData.csv
   或指定某天（YYYY-MM-DD）：
   python wim_analysis_materials.py WIMData.csv 2024-03-05

输出：
- figs/ 下保存所有图片
- summary.txt 保存关键统计结果
- （可选）如果安装 python-docx，会生成 report.docx

注意：
- 绝对不要用 "._WIMData.csv"（macOS 生成的元数据文件，二进制，读会报 UnicodeDecodeError）。
"""

import csv
import os
import sys
from collections import Counter
from datetime import datetime, date

import matplotlib.pyplot as plt

# ------------------------ 基础配置 ------------------------

VEHICLE_NAME = {
    "A": "小客车",
    "B": "小型货车",
    "C": "大客车",
    "D": "中型货车",
    "E": "大型货车",
    "F": "特大型货车",
    "G": "铰接或拖挂汽车",
    "I": "不参与统计",
}

VALID_TYPES = ["A", "B", "C", "D", "E", "F", "G"]  # I 不统计


def detect_encoding(file_path: str) -> str:
    """用 chardet 猜测文本编码（讲义第49页），失败则回退 utf-8。"""
    try:
        import chardet
    except ImportError:
        return "utf-8"

    with open(file_path, "rb") as f:
        raw = f.read(200000)  # 读少量字节即可
    result = chardet.detect(raw)
    enc = result.get("encoding") or "utf-8"
    return enc


def safe_int(x, default=None):
    try:
        return int(x)
    except Exception:
        return default


def safe_float(x, default=None):
    try:
        return float(x)
    except Exception:
        return default


def parse_time(pass_time_str: str):
    # PassTime 格式：YYYY-MM-DD HH:MM:SS
    try:
        return datetime.strptime(pass_time_str, "%Y-%m-%d %H:%M:%S")
    except Exception:
        return None


def ensure_real_wim_csv(file_path: str) -> None:
    """检查是否误用了 macOS 的 ._ 文件，或文件太小。"""
    base = os.path.basename(file_path)
    if base.startswith("._"):
        raise ValueError(
            "你选择的是 '._' 开头的文件（macOS 元数据文件），不是 WIMData.csv 正式数据！\n"
            "请改用 WIMData.csv（不带 '._'，且文件大小通常几十 MB 以上）。"
        )
    if os.path.exists(file_path):
        size = os.path.getsize(file_path)
        if size < 5 * 1024 * 1024:
            # 5MB 以下基本不可能是正式数据
            print(f"[警告] 该文件大小只有 {size} 字节，可能不是正式 WIMData.csv（通常很大）。")


def plot_pie(counter: Counter, title: str, out_path: str) -> None:
    labels = []
    sizes = []
    for k in VALID_TYPES:
        if counter.get(k, 0) > 0:
            labels.append(f"{k}({VEHICLE_NAME[k]})")
            sizes.append(counter[k])

    if not sizes:
        print(f"[跳过] {title} 没有可绘制数据")
        return

    fig, ax = plt.subplots(figsize=(7, 7))
    ax.pie(sizes, labels=labels, autopct="%1.1f%%")  # 讲义第90页
    ax.set_title(title)
    plt.tight_layout()
    plt.savefig(out_path, dpi=200)
    plt.close(fig)


def plot_lane_weight_stats(lane_stats: dict, out_path: str) -> None:
    lanes = sorted(lane_stats.keys())
    means, mins_, maxs_ = [], [], []

    for lane in lanes:
        s = lane_stats[lane]
        mean = s["sum"] / s["count"] if s["count"] else 0
        means.append(mean)
        mins_.append(s["min"])
        maxs_.append(s["max"])

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(lanes, means, marker="o", label="Mean")
    ax.plot(lanes, mins_, marker="o", label="Min")
    ax.plot(lanes, maxs_, marker="o", label="Max")
    ax.set_xlabel("Lane")
    ax.set_ylabel("TotalWeight")
    ax.set_title("Lane TotalWeight: Mean/Min/Max")
    ax.grid(True)
    ax.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=200)
    plt.close(fig)


def plot_hour_flow(hour_counts, title: str, out_path: str) -> None:
    hours = list(range(24))
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(hours, hour_counts, marker="o")
    ax.set_xlabel("Hour")
    ax.set_ylabel("Vehicle Count")
    ax.set_title(title)
    ax.grid(True)
    plt.tight_layout()
    plt.savefig(out_path, dpi=200)
    plt.close(fig)


def plot_weekday_weekend(weekday_counts, weekend_counts, out_path: str) -> None:
    hours = list(range(24))
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(hours, weekday_counts, marker="o", label="Weekday (Mon-Fri)")
    ax.plot(hours, weekend_counts, marker="o", label="Weekend (Sat-Sun)")
    ax.set_xlabel("Hour")
    ax.set_ylabel("Vehicle Count")
    ax.set_title("Hourly Flow: Weekday vs Weekend")
    ax.grid(True)
    ax.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=200)
    plt.close(fig)


def plot_weekday_weekend_by_type(weekday_type, weekend_type, out_dir: str) -> None:
    hours = list(range(24))
    for vt in VALID_TYPES:
        if vt not in weekday_type and vt not in weekend_type:
            continue
        wd = weekday_type.get(vt, [0] * 24)
        we = weekend_type.get(vt, [0] * 24)
        fig, ax = plt.subplots(figsize=(10, 5))
        ax.plot(hours, wd, marker="o", label=f"Weekday {vt}({VEHICLE_NAME[vt]})")
        ax.plot(hours, we, marker="o", label=f"Weekend {vt}({VEHICLE_NAME[vt]})")
        ax.set_xlabel("Hour")
        ax.set_ylabel("Vehicle Count")
        ax.set_title(f"Hourly Flow by Type: {vt}({VEHICLE_NAME[vt]})")
        ax.grid(True)
        ax.legend()
        plt.tight_layout()
        out_path = os.path.join(out_dir, f"07_type_{vt}_weekday_vs_weekend.png")
        plt.savefig(out_path, dpi=200)
        plt.close(fig)


def is_full_day(hour_set: set) -> bool:
    return len(hour_set) == 24


def generate_docx(summary_txt_path: str, figs_dir: str, out_docx_path: str) -> bool:
    """可选：生成 Word 报告（需要 python-docx）。

    讲义里没有讲 python-docx，所以这里做成“可选”。
    """
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception:
        return False

    doc = Document()
    doc.add_heading("WIM 数据分析报告", level=0)

    doc.add_heading("关键统计结果", level=1)
    with open(summary_txt_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("图表", level=1)

    # 按文件名顺序插图
    for name in sorted(os.listdir(figs_dir)):
        if not name.lower().endswith(".png"):
            continue
        doc.add_paragraph(name)
        doc.add_picture(os.path.join(figs_dir, name), width=Inches(6.2))

    doc.save(out_docx_path)
    return True


def main():
    # ------------------------ 读取参数 ------------------------
    csv_path = sys.argv[1] if len(sys.argv) >= 2 else "WIMData.csv"
    day_str = sys.argv[2] if len(sys.argv) >= 3 else None

    ensure_real_wim_csv(csv_path)

    if not os.path.exists(csv_path):
        print(f"找不到文件：{csv_path}")
        print("请确认 WIMData.csv 与脚本在同一目录，或在命令行传入文件路径。")
        return

    # 输出目录
    out_dir = "output"
    figs_dir = os.path.join(out_dir, "figs")
    os.makedirs(figs_dir, exist_ok=True)

    # 编码探测（讲义第49页），并显式指定编码（讲义第53页最佳实践）
    encoding = detect_encoding(csv_path)
    print(f"检测到编码：{encoding}")

    # ------------------------ 统计容器 ------------------------
    overall_type = Counter()
    lane_type = {}  # lane -> Counter

    # 方向统计：1~3 为方向1，4~6 为方向2
    dir_total = {"Dir_1(1-3)": 0, "Dir_2(4-6)": 0}
    dir_type = {"Dir_1(1-3)": Counter(), "Dir_2(4-6)": Counter()}

    # 车道重量统计
    lane_weight = {}  # lane -> {sum,count,min,max}

    # 用于找“完整的一天”（24小时都有）
    day_hours_seen = {}  # day -> set(hours)
    day_hour_counts = {}  # day -> [24]

    # 选做：工作日 vs 双休日
    weekday_hour = [0] * 24
    weekend_hour = [0] * 24
    weekday_type_hour = {}  # vt -> [24]
    weekend_type_hour = {}

    # ------------------------ 读 CSV 并统计 ------------------------
    # 讲义第52页：csv.DictReader
    with open(csv_path, "r", newline="", encoding=encoding) as f:
        reader = csv.DictReader(f)

        # 必需字段检查
        need_cols = ["PassTime", "Lane", "TotalWeight", "VehicleType"]
        for c in need_cols:
            if c not in reader.fieldnames:
                raise ValueError(f"CSV 缺少字段：{c}，请确认文件是否为 WIMData.csv")

        row_count = 0
        for row in reader:
            row_count += 1

            vt = (row.get("VehicleType") or "").strip()
            if vt not in VALID_TYPES:
                # I 或空值：不参与统计
                continue

            lane = safe_int(row.get("Lane"), default=None)
            tw = safe_float(row.get("TotalWeight"), default=None)
            t = parse_time((row.get("PassTime") or "").strip())

            # 基础过滤
            if lane is None or tw is None or t is None:
                continue

            # 题1：总体车型
            overall_type[vt] += 1

            # 题2：每车道车型
            if lane not in lane_type:
                lane_type[lane] = Counter()
            lane_type[lane][vt] += 1

            # 题3：方向统计
            if 1 <= lane <= 3:
                dkey = "Dir_1(1-3)"
            elif 4 <= lane <= 6:
                dkey = "Dir_2(4-6)"
            else:
                dkey = None

            if dkey is not None:
                dir_total[dkey] += 1
                dir_type[dkey][vt] += 1

            # 题4：车道重量统计
            if lane not in lane_weight:
                lane_weight[lane] = {"sum": 0.0, "count": 0, "min": tw, "max": tw}
            lane_weight[lane]["sum"] += tw
            lane_weight[lane]["count"] += 1
            if tw < lane_weight[lane]["min"]:
                lane_weight[lane]["min"] = tw
            if tw > lane_weight[lane]["max"]:
                lane_weight[lane]["max"] = tw

            # 题5：按天按小时
            day = t.date()
            hour = t.hour

            if day not in day_hours_seen:
                day_hours_seen[day] = set()
                day_hour_counts[day] = [0] * 24

            day_hours_seen[day].add(hour)
            day_hour_counts[day][hour] += 1

            # 选做：工作日/周末
            # weekday(): Monday=0 ... Sunday=6
            is_weekend = (day.weekday() >= 5)
            if is_weekend:
                weekend_hour[hour] += 1
            else:
                weekday_hour[hour] += 1

            # 分车型
            if is_weekend:
                if vt not in weekend_type_hour:
                    weekend_type_hour[vt] = [0] * 24
                weekend_type_hour[vt][hour] += 1
            else:
                if vt not in weekday_type_hour:
                    weekday_type_hour[vt] = [0] * 24
                weekday_type_hour[vt][hour] += 1

            # 进度提示
            if row_count % 200000 == 0:
                print(f"已处理 {row_count} 行...")

    print(f"完成读取与统计，总处理行数（含过滤前）：{row_count}")

    # ------------------------ 选择完整日期（题5） ------------------------
    chosen_day = None

    # 如果用户指定日期，优先用
    if day_str:
        try:
            chosen_day = datetime.strptime(day_str, "%Y-%m-%d").date()
        except Exception:
            print("日期参数格式应为 YYYY-MM-DD，例如 2024-03-05。将自动选择完整日期。")
            chosen_day = None

    # 否则自动找 2024-03-05 ~ 2024-03-28 内的“完整一天”
    if chosen_day is None:
        start = date(2024, 3, 5)
        end = date(2024, 3, 28)
        cand = []
        d = start
        while d <= end:
            if d in day_hours_seen and is_full_day(day_hours_seen[d]):
                cand.append(d)
            d = date.fromordinal(d.toordinal() + 1)

        if cand:
            chosen_day = cand[0]  # 取最早的完整日
        else:
            # 兜底：选小时覆盖最多的一天
            best_day = None
            best_cnt = -1
            for d, hs in day_hours_seen.items():
                if len(hs) > best_cnt:
                    best_cnt = len(hs)
                    best_day = d
            chosen_day = best_day

    # ------------------------ 输出 summary.txt ------------------------
    summary_path = os.path.join(out_dir, "summary.txt")
    with open(summary_path, "w", encoding="utf-8") as out:
        out.write("=== 题1：总体车型占比（不含 I）===\n")
        total_valid = sum(overall_type.values())
        out.write(f"总车辆数（有效车型）：{total_valid}\n")
        for vt in VALID_TYPES:
            c = overall_type.get(vt, 0)
            if total_valid:
                out.write(f"{vt}({VEHICLE_NAME[vt]}): {c}  占比 {c/total_valid*100:.2f}%\n")
        out.write("\n")

        out.write("=== 题3：两个方向总数与分车型 ===\n")
        for dkey in ["Dir_1(1-3)", "Dir_2(4-6)"]:
            out.write(f"{dkey} 总数：{dir_total[dkey]}\n")
            for vt in VALID_TYPES:
                out.write(f"  {vt}({VEHICLE_NAME[vt]}): {dir_type[dkey].get(vt,0)}\n")
        out.write("\n")

        out.write("=== 题4：各车道 TotalWeight 统计（均值/最小/最大）===\n")
        for lane in sorted(lane_weight.keys()):
            s = lane_weight[lane]
            mean = s["sum"] / s["count"] if s["count"] else 0
            out.write(
                f"Lane {lane}: mean={mean:.2f}, min={s['min']:.2f}, max={s['max']:.2f}, n={s['count']}\n"
            )
        out.write("\n")

        out.write("=== 题5：选取日期逐小时车流量 ===\n")
        out.write(f"选取日期：{chosen_day}\n")
        if chosen_day in day_hour_counts:
            counts = day_hour_counts[chosen_day]
            for h in range(24):
                out.write(f"{h:02d}:00 - {counts[h]}\n")
        else:
            out.write("（该日期无数据）\n")

    print(f"已写入：{summary_path}")

    # ------------------------ 画图（题1/2/4/5/7） ------------------------

    # 题1：总体饼图
    plot_pie(
        overall_type,
        title="Overall Vehicle Type Share",
        out_path=os.path.join(figs_dir, "01_overall_vehicle_type_pie.png"),
    )

    # 题2：每车道饼图
    for lane in sorted(lane_type.keys()):
        plot_pie(
            lane_type[lane],
            title=f"Lane {lane} Vehicle Type Share",
            out_path=os.path.join(figs_dir, f"02_lane_{lane}_vehicle_type_pie.png"),
        )

    # 题4：重量统计折线图
    plot_lane_weight_stats(
        lane_weight,
        out_path=os.path.join(figs_dir, "04_lane_totalweight_stats.png"),
    )

    # 题5：某天逐小时车流量
    if chosen_day in day_hour_counts:
        plot_hour_flow(
            day_hour_counts[chosen_day],
            title=f"Hourly Flow on {chosen_day}",
            out_path=os.path.join(figs_dir, "05_hourly_flow_selected_day.png"),
        )

    # 题7（选做）：工作日 vs 周末
    plot_weekday_weekend(
        weekday_hour,
        weekend_hour,
        out_path=os.path.join(figs_dir, "06_weekday_vs_weekend_hourly.png"),
    )
    plot_weekday_weekend_by_type(weekday_type_hour, weekend_type_hour, figs_dir)

    print(f"图片已保存到：{figs_dir}")

    # ------------------------ 可选：生成 Word 报告 ------------------------
    docx_path = os.path.join(out_dir, "report.docx")
    ok = generate_docx(summary_path, figs_dir, docx_path)
    if ok:
        print(f"已生成 Word 报告：{docx_path}")
    else:
        print("[提示] 未生成 report.docx（可能没安装 python-docx）。你可以：pip install python-docx")


if __name__ == "__main__":
    main()
